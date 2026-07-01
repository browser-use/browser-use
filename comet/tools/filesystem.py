"""
Comet Tools — Excel / Word / PDF / CSV / JSON
Full file-system layer on top of browser-use.
"""
from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.shared import Pt
import pdfplumber

from comet.utils.logger import CometLogger


class FileSystemTools:
    """Read and write Excel, Word, PDF, CSV and JSON files."""

    def __init__(self, logger: CometLogger):
        self.logger = logger

    # ── Excel ──────────────────────────────────────────────────

    def read_excel_data(self, file_path: str,
                        sheet_name: Any = 0,
                        max_rows: int = 1000) -> str:
        try:
            p = Path(file_path)
            if not p.exists():
                return f"❌ Fichier introuvable : {file_path}"
            df = pd.read_excel(str(p), sheet_name=sheet_name,
                               nrows=max_rows, engine="openpyxl")
            df.columns = [str(c).strip() for c in df.columns]
            df = df.dropna(how="all")
            self.logger.success(f"Excel lu : {p.name} — {len(df)} lignes")
            return (
                f"✅ {p.name} | colonnes: {list(df.columns)} | "
                f"{len(df)} lignes\n\n" + df.to_string(index=False, max_rows=50)
            )
        except Exception as e:
            return f"❌ read_excel_data : {e}"

    def write_to_excel(self, data: list[dict], file_path: str,
                       sheet_name: str = "Comet") -> str:
        try:
            p = Path(file_path)
            p.parent.mkdir(parents=True, exist_ok=True)
            df = pd.DataFrame(data)
            with pd.ExcelWriter(str(p), engine="openpyxl") as writer:
                df.to_excel(writer, index=False, sheet_name=sheet_name)
                ws = writer.sheets[sheet_name]
                for col in ws.columns:
                    max_len = max(len(str(c.value or "")) for c in col)
                    ws.column_dimensions[col[0].column_letter].width = min(
                        max_len + 4, 50)
            self.logger.success(
                f"Excel sauvegardé : {p.name} ({len(data)} lignes)")
            return f"✅ {p.name} — {len(data)} lignes écrites"
        except Exception as e:
            return f"❌ write_to_excel : {e}"

    def append_to_excel(self, new_data: list[dict],
                        file_path: str) -> str:
        try:
            p = Path(file_path)
            if p.exists():
                existing = pd.read_excel(str(p), engine="openpyxl")
                combined = pd.concat(
                    [existing, pd.DataFrame(new_data)], ignore_index=True)
            else:
                combined = pd.DataFrame(new_data)
            combined.to_excel(str(p), index=False, engine="openpyxl")
            return f"✅ {len(new_data)} lignes ajoutées — total {len(combined)}"
        except Exception as e:
            return f"❌ append_to_excel : {e}"

    # ── CSV ────────────────────────────────────────────────────

    def read_csv(self, file_path: str, delimiter: str = ",") -> str:
        try:
            df = pd.read_csv(file_path, delimiter=delimiter)
            return (f"✅ CSV : {Path(file_path).name} | "
                    f"{len(df)} lignes | {list(df.columns)}\n\n"
                    + df.to_string(index=False, max_rows=30))
        except Exception as e:
            return f"❌ read_csv : {e}"

    def write_csv(self, data: list[dict], file_path: str) -> str:
        try:
            p = Path(file_path)
            p.parent.mkdir(parents=True, exist_ok=True)
            with open(str(p), "w", newline="", encoding="utf-8") as f:
                writer = csv.DictWriter(f, fieldnames=data[0].keys())
                writer.writeheader()
                writer.writerows(data)
            return f"✅ CSV : {p.name} ({len(data)} lignes)"
        except Exception as e:
            return f"❌ write_csv : {e}"

    # ── Word ───────────────────────────────────────────────────

    def generate_word_document(self, text_content: str,
                                file_path: str,
                                title: str = "Rapport Comet") -> str:
        try:
            p = Path(file_path)
            p.parent.mkdir(parents=True, exist_ok=True)
            doc = Document()
            h = doc.add_heading(title, level=1)
            h.style.font.size = Pt(18)
            for block in text_content.split("\n\n"):
                block = block.strip()
                if not block:
                    continue
                if block.startswith("# "):
                    doc.add_heading(block[2:], level=2)
                elif block.startswith("## "):
                    doc.add_heading(block[3:], level=3)
                elif block.startswith("- ") or block.startswith("* "):
                    for item in block.split("\n"):
                        doc.add_paragraph(
                            item.lstrip("-* ").strip(),
                            style="List Bullet")
                else:
                    para = doc.add_paragraph(block)
                    para.style.font.size = Pt(11)
            doc.save(str(p))
            self.logger.success(f"Word généré : {p.name}")
            return f"✅ {p.name} créé"
        except Exception as e:
            return f"❌ generate_word_document : {e}"

    # ── PDF ────────────────────────────────────────────────────

    def read_pdf(self, file_path: str, max_pages: int = 20) -> str:
        try:
            p = Path(file_path)
            if not p.exists():
                return f"❌ PDF introuvable : {file_path}"
            parts = []
            with pdfplumber.open(str(p)) as pdf:
                for i, page in enumerate(pdf.pages[:max_pages], 1):
                    parts.append(f"--- Page {i} ---\n{page.extract_text() or ''}")
            return f"✅ PDF : {p.name}\n\n" + "\n\n".join(parts)[:5000]
        except Exception as e:
            return f"❌ read_pdf : {e}"

    # ── JSON ───────────────────────────────────────────────────

    def read_json(self, file_path: str) -> str:
        try:
            data = json.loads(Path(file_path).read_text(encoding="utf-8"))
            return ("✅ JSON lu\n"
                    + json.dumps(data, ensure_ascii=False, indent=2)[:3000])
        except Exception as e:
            return f"❌ read_json : {e}"

    def write_json(self, data: Any, file_path: str) -> str:
        try:
            p = Path(file_path)
            p.parent.mkdir(parents=True, exist_ok=True)
            p.write_text(
                json.dumps(data, ensure_ascii=False, indent=2),
                encoding="utf-8")
            return f"✅ JSON : {p.name}"
        except Exception as e:
            return f"❌ write_json : {e}"
